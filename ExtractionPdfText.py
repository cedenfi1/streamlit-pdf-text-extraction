import streamlit as st
import PyPDF2
from docx import Document
from io import BytesIO

st.title("PDF to Word Document Converter")
st.write("### Step 1: Upload your PDF files")
st.write("Upload as many PDF files as needed.")

# Allow multiple PDF file uploads
uploaded_files = st.file_uploader("Choose PDF files", type="pdf", accept_multiple_files=True)

if uploaded_files:
    st.write(f"You've uploaded {len(uploaded_files)} file(s).")
    
    st.write("### Step 2: Extract text and create Word document")
    if st.button("Generate Word Document"):
        # Create a new Word document
        document = Document()
        
        # Process each uploaded PDF file
        for idx, uploaded_file in enumerate(uploaded_files, start=1):
            st.write(f"Processing {uploaded_file.name}...")
            try:
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                pdf_text = ""
                # Extract text from each page in the PDF
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text() or ""
                    pdf_text += page_text + "\n"
                
                # Add a title or header for the current PDF in the doc
                document.add_heading(f"Content from {uploaded_file.name}", level=2)
                # Add the extracted text from the PDF
                document.add_paragraph(pdf_text)
                # Add a horizontal rule (simulate with a paragraph of underscores) and extra breaks
                document.add_paragraph("\n" + "_" * 50 + "\n")
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {e}")
        
        # Save the document to a bytes buffer
        doc_buffer = BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)
        
        st.success("Word document created successfully!")
        st.download_button(
            label="Download Word Document",
            data=doc_buffer,
            file_name="combined_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
